from docx import Document
from docx.oxml.ns import qn
doc = Document("reports/reporte_2026_03.docx")
fonts = set()
for p in doc.paragraphs:
    for r in p.runs:
        if r.font.name: fonts.add(r.font.name)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.font.name: fonts.add(r.font.name)
print(f"FONTS: {fonts}")
print("\nTABLE HEADERS:")
for i, t in enumerate(doc.tables):
    hdrs = [c.text.strip() for c in t.rows[0].cells]
    print(f"  T{i}: {hdrs}")
for p in doc.paragraphs:
    if "uptime promedio" in p.text.lower():
        print(f"\nUPTIME: {p.text[:150]}")
bms = [e.get(qn("w:name")) for e in doc.element.iter() if e.tag == qn("w:bookmarkStart") and e.get(qn("w:name")) != "_GoBack"]
print(f"\nBOOKMARKS ({len(bms)}): {bms[:8]}...")
hls = [e.get(qn("w:anchor")) for e in doc.element.iter() if e.tag == qn("w:hyperlink") and e.get(qn("w:anchor"))]
print(f"HYPERLINKS ({len(hls)}): {hls[:8]}...")
for p in doc.paragraphs:
    if "ndice" in p.text.lower():
        print(f"\nTOC: '{p.text[:80]}'")
        break
